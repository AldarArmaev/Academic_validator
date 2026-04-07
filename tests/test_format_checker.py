# tests/test_format_checker.py
import pytest
import sys
import uuid
from pathlib import Path

# Добавляем backend в путь импорта
sys.path.insert(0, str(Path(__file__).parent.parent / "normocontroller" / "backend"))

from docx import Document
from schemas import ReportError, ValidationReport
from contour_a.format_checker import (
    check_font_formatting,
    check_paragraph_formatting,
    check_required_sections,
    check_page_margins,
    validate_format
)


class TestCheckFontFormatting:
    def test_correct_font_no_errors(self, fixture_path):
        """Проверка отсутствия ошибок при правильном шрифте."""
        doc = Document(str(fixture_path("correct.docx")))
        errors = check_font_formatting(doc)
        assert len(errors) == 0

    def test_wrong_font_detected(self, fixture_path):
        """Проверка обнаружения неверного шрифта."""
        doc = Document(str(fixture_path("wrong_font.docx")))
        errors = check_font_formatting(doc)
        assert len(errors) > 0

    def test_wrong_font_error_type(self, fixture_path):
        """Проверка типа ошибки для неверного шрифта."""
        doc = Document(str(fixture_path("wrong_font.docx")))
        errors = check_font_formatting(doc)
        if errors:
            assert errors[0].type == "formatting"

    def test_wrong_font_severity(self, fixture_path):
        """Проверка серьёзности ошибки для неверного шрифта."""
        doc = Document(str(fixture_path("wrong_font.docx")))
        errors = check_font_formatting(doc)
        if errors:
            assert errors[0].severity == "error"

    def test_wrong_font_has_recommendation(self, fixture_path):
        """Проверка наличия рекомендации для ошибки шрифта."""
        doc = Document(str(fixture_path("wrong_font.docx")))
        errors = check_font_formatting(doc)
        if errors:
            assert len(errors[0].recommendation) > 0

    def test_wrong_font_fragment_not_empty(self, fixture_path):
        """Проверка, что фрагмент ошибки не пустой."""
        doc = Document(str(fixture_path("wrong_font.docx")))
        errors = check_font_formatting(doc)
        if errors:
            assert len(errors[0].fragment) > 0

    def test_wrong_font_expected_value(self, fixture_path):
        """Проверка ожидаемого значения для ошибки шрифта."""
        doc = Document(str(fixture_path("wrong_font.docx")))
        errors = check_font_formatting(doc)
        if errors:
            assert errors[0].expected_value == "Times New Roman"


class TestCheckParagraphFormatting:
    def test_correct_indent_no_errors(self, fixture_path):
        """Проверка отсутствия ошибок при правильных отступах."""
        doc = Document(str(fixture_path("correct.docx")))
        errors = check_paragraph_formatting(doc)
        assert len(errors) == 0

    def test_wrong_indent_detected(self, fixture_path):
        """Проверка обнаружения неверных отступов."""
        doc = Document(str(fixture_path("wrong_indent.docx")))
        errors = check_paragraph_formatting(doc)
        assert len(errors) > 0

    def test_wrong_spacing_detected(self, fixture_path):
        """Проверка обнаружения неверных интервалов."""
        doc = Document(str(fixture_path("wrong_spacing.docx")))
        errors = check_paragraph_formatting(doc)
        assert len(errors) > 0

    def test_wrong_spacing_expected_value(self, fixture_path):
        """Проверка ожидаемого значения интервала."""
        doc = Document(str(fixture_path("wrong_spacing.docx")))
        errors = check_paragraph_formatting(doc)
        if errors:
            assert errors[0].expected_value == "360"

    def test_wrong_alignment_detected(self, fixture_path):
        """Проверка обнаружения неверного выравнивания."""
        doc = Document(str(fixture_path("wrong_alignment.docx")))
        errors = check_paragraph_formatting(doc)
        assert len(errors) > 0

    def test_wrong_alignment_expected_value(self, fixture_path):
        """Проверка ожидаемого значения выравнивания."""
        doc = Document(str(fixture_path("wrong_alignment.docx")))
        errors = check_paragraph_formatting(doc)
        if errors:
            assert errors[0].expected_value == "both"

    def test_paragraph_error_has_location(self, fixture_path):
        """Проверка наличия местоположения у ошибки форматирования абзаца."""
        doc = Document(str(fixture_path("wrong_indent.docx")))
        errors = check_paragraph_formatting(doc)
        if errors:
            assert hasattr(errors[0], 'location')
            assert errors[0].location.page >= 0


class TestCheckRequiredSections:
    def test_all_sections_present(self, fixture_path):
        """Проверка документа со всеми разделами."""
        doc = Document(str(fixture_path("correct.docx")))
        errors = check_required_sections(doc)
        assert len(errors) == 0

    def test_missing_section_detected(self, fixture_path):
        """Проверка обнаружения отсутствующего раздела."""
        doc = Document(str(fixture_path("missing_sections.docx")))
        errors = check_required_sections(doc)
        assert len(errors) > 0

    def test_missing_section_type(self, fixture_path):
        """Проверка типа ошибки для отсутствующего раздела."""
        doc = Document(str(fixture_path("missing_sections.docx")))
        errors = check_required_sections(doc)
        if errors:
            assert errors[0].type == "formatting"

    def test_missing_section_has_rule(self, fixture_path):
        """Проверка наличия описания правила для отсутствующего раздела."""
        doc = Document(str(fixture_path("missing_sections.docx")))
        errors = check_required_sections(doc)
        if errors:
            assert len(errors[0].rule) > 0

    def test_missing_introduction(self, fixture_path):
        """Проверка обнаружения отсутствия раздела 'Введение'."""
        doc = Document(str(fixture_path("missing_sections.docx")))
        errors = check_required_sections(doc)
        # Проверяем, что среди ошибок есть отсутствие введения
        intro_errors = [e for e in errors if "введен" in e.expected_value.lower() or "введен" in e.rule.lower()]
        assert len(intro_errors) > 0

    def test_missing_conclusion(self, fixture_path):
        """Проверка обнаружения отсутствия раздела 'Заключение'."""
        doc = Document(str(fixture_path("missing_sections.docx")))
        errors = check_required_sections(doc)
        # Проверяем, что среди ошибок есть отсутствие заключения
        conclusion_errors = [e for e in errors if "заключен" in e.expected_value.lower() or "заключен" in e.rule.lower()]
        assert len(conclusion_errors) > 0

    def test_missing_references(self, fixture_path):
        """Проверка обнаружения отсутствия 'Списка литературы'."""
        doc = Document(str(fixture_path("missing_sections.docx")))
        errors = check_required_sections(doc)
        # Проверяем, что среди ошибок есть отсутствие списка литературы
        ref_errors = [e for e in errors if "список" in e.expected_value.lower() or "литератур" in e.expected_value.lower() or "список" in e.rule.lower()]
        assert len(ref_errors) > 0


class TestCheckPageMargins:
    def test_correct_margins_no_errors(self, fixture_path):
        """Проверка отсутствия ошибок при правильных полях."""
        doc = Document(str(fixture_path("correct.docx")))
        errors = check_page_margins(doc)
        assert len(errors) == 0

    def test_wrong_margins_detected(self, fixture_path):
        """Проверка обнаружения неверных полей."""
        doc = Document(str(fixture_path("wrong_margins.docx")))
        errors = check_page_margins(doc)
        assert len(errors) > 0

    def test_wrong_margin_error_type(self, fixture_path):
        """Проверка типа ошибки для неверных полей."""
        doc = Document(str(fixture_path("wrong_margins.docx")))
        errors = check_page_margins(doc)
        if errors:
            assert errors[0].type == "formatting"

    def test_wrong_margin_found_value_numeric(self, fixture_path):
        """Проверка, что найденное значение поля числовое."""
        doc = Document(str(fixture_path("wrong_margins.docx")))
        errors = check_page_margins(doc)
        if errors:
            # found_value должно быть преобразуемо в int
            try:
                int(errors[0].found_value)
            except ValueError:
                pytest.fail("found_value должно быть числовым")

    def test_wrong_margin_expected_value(self, fixture_path):
        """Проверка ожидаемого значения для ошибки поля (1701 твип)."""
        doc = Document(str(fixture_path("wrong_margins.docx")))
        errors = check_page_margins(doc)
        if errors:
            assert errors[0].expected_value == "1701"


class TestValidateFormat:
    def test_validate_format_returns_report(self, fixture_path):
        """Проверка, что validate_format возвращает отчёт."""
        report = validate_format(str(fixture_path("correct.docx")))
        assert isinstance(report, ValidationReport)

    def test_validate_format_has_doc_id(self, fixture_path):
        """Проверка наличия doc_id в отчёте."""
        report = validate_format(str(fixture_path("correct.docx")))
        assert hasattr(report, 'doc_id')
        assert len(report.doc_id) > 0
        # Дополнительная проверка: doc_id должен быть валидным UUID
        uuid.UUID(report.doc_id)

    def test_validate_format_doc_id_is_uuid(self, fixture_path):
        """Проверка, что doc_id является UUID."""
        report = validate_format(str(fixture_path("correct.docx")))
        uuid.UUID(report.doc_id)  # Не должно вызывать исключений

    def test_validate_format_has_summary(self, fixture_path):
        """Проверка наличия сводки в отчёте."""
        report = validate_format(str(fixture_path("correct.docx")))
        assert hasattr(report, 'summary')
        assert report.summary.total_errors >= 0

    def test_validate_format_invalid_path(self, fixture_path):
        """Проверка обработки неверного пути."""
        with pytest.raises(Exception):
            validate_format("/nonexistent/path/file.docx")

    def test_validate_format_counts_errors(self, fixture_path):
        """Проверка подсчёта ошибок в отчёте."""
        # Примечание: wrong_font.docx может содержать и другие ошибки форматирования
        report = validate_format(str(fixture_path("wrong_font.docx")))
        assert report.summary.total_errors == len(report.errors)
        assert report.summary.formatting + report.summary.style + report.summary.citations == report.summary.total_errors
