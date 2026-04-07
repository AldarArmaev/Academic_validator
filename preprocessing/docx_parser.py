"""Module for parsing DOCX files."""

import math
from typing import Dict, List

try:
    from docx import Document
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

try:
    from mathml2latex import convert as mathml_to_latex
    MATHML_AVAILABLE = True
except ImportError:
    MATHML_AVAILABLE = False


def is_inside_table(paragraph):
    """Check if paragraph is inside a table."""
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith('tbl'):
            return True
        parent = parent.getparent()
    return False


def has_drawing(paragraph):
    """Check if paragraph contains a drawing/image."""
    return 'drawing' in paragraph._element.xml


def parse_docx(file_path: str) -> Dict:
    """
    Extract paragraphs and metadata from a DOCX file.
    
    Args:
        file_path: Path to the .docx file.
        
    Returns:
        Dict with keys:
            - paragraphs: list of dicts with text, style, runs_count
            - metadata: dict with total_paragraphs, total_pages
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise FileNotFoundError(f"Cannot open file: {file_path}") from e
    
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append({
            "text": para.text,
            "style": para.style.name if para.style else "Normal",
            "runs_count": len(para.runs)
        })
    
    total_paragraphs = len(paragraphs)
    # Estimate pages as ceil(paragraphs / 40). 
    # Note: This is an approximation for MVP; accurate page count requires 
    # analyzing section breaks and layout metrics which docx doesn't expose directly.
    total_pages = math.ceil(total_paragraphs / 40) if total_paragraphs > 0 else 0
    
    return {
        "paragraphs": paragraphs,
        "metadata": {
            "total_paragraphs": total_paragraphs,
            "total_pages": total_pages
        }
    }


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file as a single string.
    
    Args:
        file_path: Path to the .docx file.
        
    Returns:
        Full text content of the document.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise FileNotFoundError(f"Cannot open file: {file_path}") from e
    
    texts = [para.text for para in doc.paragraphs]
    return "\n".join(texts)


def convert_to_markdown(file_path: str) -> str:
    """
    Convert a DOCX file to Markdown format.
    
    Args:
        file_path: Path to the .docx file.
        
    Returns:
        Markdown string representation of the document.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise FileNotFoundError(f"Cannot open file: {file_path}") from e
    
    markdown_lines = []
    
    for para in doc.paragraphs:
        # Skip paragraphs inside tables
        if is_inside_table(para):
            continue
        
        # Skip paragraphs with drawings/images
        if has_drawing(para):
            continue
        
        text = para.text
        style_name = para.style.name if para.style else "Normal"
        
        # Handle headings
        if style_name == "Heading 1":
            markdown_lines.append(f"# {text}")
        elif style_name == "Heading 2":
            markdown_lines.append(f"## {text}")
        elif style_name == "Heading 3":
            markdown_lines.append(f"### {text}")
        else:
            # Check for OMML formulas in runs
            has_formula = False
            formula_text = text
            
            if MATHML_AVAILABLE:
                for run in para.runs:
                    # Try to detect and convert OMML formulas
                    # OMML is typically in XML format within the run
                    try:
                        # Access the run's XML element
                        run_xml = run._element.xml
                        if 'm:' in run_xml or '<m:' in run_xml:
                            # Attempt to extract and convert MathML
                            # This is a simplified approach
                            has_formula = True
                    except Exception:
                        pass
            
            if has_formula and MATHML_AVAILABLE:
                # If formula detected and converter available, attempt conversion
                # Note: Full OMML to LaTeX conversion is complex
                # This is a placeholder for actual conversion logic
                pass
            
            markdown_lines.append(text)
    
    return "\n".join(markdown_lines)
