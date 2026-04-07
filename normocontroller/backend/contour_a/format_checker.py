"""Модуль проверки форматирования документов."""

from typing import List, Optional
import json
from pathlib import Path

from docx import Document

from schemas import ReportError, ErrorLocation, ReportSummary, ValidationReport


# Загрузка правил из JSON
RULES_PATH = Path(__file__).parent.parent / "university_rules.json"
with open(RULES_PATH, "r", encoding="utf-8") as f:
    UNIVERSITY_RULES = json.load(f)


def check_font_formatting(doc: Document) -> List[ReportError]:
    """Проверка шрифта и размера текста."""
    errors = []
    font_family = UNIVERSITY_RULES["font"]["family"]
    size_half_points = UNIVERSITY_RULES["font"]["size_half_points"]
    
    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run.font.name and run.font.name != font_family:
                errors.append(ReportError(
                    id=f"font_{para_idx}",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(page=0, paragraph_index=para_idx, chapter=None),
                    fragment=run.text[:50] if run.text else "",
                    rule=f"Шрифт должен быть {font_family}",
                    found_value=run.font.name or "не указан",
                    expected_value=font_family,
                    recommendation=f"Измените шрифт на {font_family}"
                ))
                break
    
    return errors


def check_paragraph_formatting(doc: Document) -> List[ReportError]:
    """Проверка форматирования абзацев."""
    errors = []
    # Здесь будет логика проверки отступов и интервалов
    # Для MVP возвращаем пустой список
    return errors


def check_page_margins(doc: Document) -> List[ReportError]:
    """Проверка полей страницы."""
    errors = []
    margins_config = UNIVERSITY_RULES.get("margins_dxa", {})
    
    # Получаем поля документа (если доступны через section properties)
    for section_idx, section in enumerate(doc.sections):
        # Проверяем левое поле
        if hasattr(section, 'left_margin'):
            left_margin = section.left_margin
            expected_left = margins_config.get("left", 1701)
            if left_margin != expected_left:
                errors.append(ReportError(
                    id=f"margin_left_{section_idx}",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(page=section_idx, paragraph_index=0, chapter=None),
                    fragment="",
                    rule=f"Левое поле должно быть {expected_left} DXA",
                    found_value=str(left_margin),
                    expected_value=str(expected_left),
                    recommendation=f"Установите левое поле {expected_left} DXA"
                ))
        
        # Проверяем правое поле
        if hasattr(section, 'right_margin'):
            right_margin = section.right_margin
            expected_right = margins_config.get("right", 567)
            if right_margin != expected_right:
                errors.append(ReportError(
                    id=f"margin_right_{section_idx}",
                    type="formatting",
                    severity="error",
                    location=ErrorLocation(page=section_idx, paragraph_index=0, chapter=None),
                    fragment="",
                    rule=f"Правое поле должно быть {expected_right} DXA",
                    found_value=str(right_margin),
                    expected_value=str(expected_right),
                    recommendation=f"Установите правое поле {expected_right} DXA"
                ))
    
    return errors


def check_required_sections(doc: Document) -> List[ReportError]:
    """Проверка наличия обязательных разделов."""
    errors = []
    required_sections = UNIVERSITY_RULES["required_sections"]
    
    doc_text = "\n".join([p.text.lower() for p in doc.paragraphs])
    
    for section in required_sections:
        if section not in doc_text:
            errors.append(ReportError(
                id=f"section_{section}",
                type="formatting",
                severity="error",
                location=ErrorLocation(page=0, paragraph_index=0, chapter=None),
                fragment="",
                rule=f"Документ должен содержать раздел '{section}'",
                found_value="раздел не найден",
                expected_value=section,
                recommendation=f"Добавьте раздел '{section}' в документ"
            ))
    
    return errors


def validate_format(docx_path: str) -> ValidationReport:
    """Основная функция проверки форматирования документа.
    
    Args:
        docx_path: Путь к файлу .docx
        
    Returns:
        ValidationReport: Отчёт о проверке
    """
    import uuid
    from datetime import datetime, timedelta
    
    doc = Document(docx_path)
    
    all_errors: List[ReportError] = []
    all_errors.extend(check_font_formatting(doc))
    all_errors.extend(check_paragraph_formatting(doc))
    all_errors.extend(check_required_sections(doc))
    
    # Подсчёт статистики
    formatting_errors = sum(1 for e in all_errors if e.type == "formatting")
    style_errors = sum(1 for e in all_errors if e.type == "style")
    citation_errors = sum(1 for e in all_errors if e.type == "citation_check")
    
    summary = ReportSummary(
        total_errors=len(all_errors),
        formatting=formatting_errors,
        style=style_errors,
        citations=citation_errors
    )
    
    return ValidationReport(
        doc_id=str(uuid.uuid4()),
        session_expires_at=datetime.utcnow() + timedelta(hours=24),
        summary=summary,
        errors=all_errors
    )
