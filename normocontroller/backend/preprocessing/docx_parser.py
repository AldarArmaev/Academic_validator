"""Парсер DOCX файлов."""

from typing import List, Dict, Any
from docx import Document


def parse_docx(file_path: str) -> Dict[str, Any]:
    """Извлекает текст и метаданные из DOCX файла.
    
    Args:
        file_path: Путь к файлу .docx
        
    Returns:
        Словарь с полями:
            - paragraphs: список абзацев
            - metadata: метаданные документа
    """
    doc = Document(file_path)
    
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append({
            "text": para.text,
            "style": para.style.name if para.style else None,
            "runs_count": len(para.runs)
        })
    
    metadata = {
        "author": doc.core_properties.author,
        "title": doc.core_properties.title,
        "subject": doc.core_properties.subject,
        "created": doc.core_properties.created,
        "modified": doc.core_properties.modified,
        "total_paragraphs": len(paragraphs),
        "total_pages": estimate_pages(doc)
    }
    
    return {
        "paragraphs": paragraphs,
        "metadata": metadata
    }


def estimate_pages(doc: Document) -> int:
    """Оценивает количество страниц в документе.
    
    Примечание: python-docx не предоставляет точное количество страниц,
    поэтому используется приблизительная оценка.
    
    Args:
        doc: Объект Document
        
    Returns:
        Приблизительное количество страниц
    """
    # Грубая оценка: ~25 строк на страницу
    total_lines = sum(len(p.text.split('\n')) for p in doc.paragraphs if p.text)
    return max(1, total_lines // 25)


def extract_text_from_docx(file_path: str) -> str:
    """Извлекает весь текст из DOCX файла.
    
    Args:
        file_path: Путь к файлу .docx
        
    Returns:
        Полный текст документа
    """
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text])
